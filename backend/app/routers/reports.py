from io import BytesIO
from typing import Any, Dict

from docx import Document
from docx.shared import Pt, Inches
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse

from app.dependencies import get_current_user
from app.services.user_service import User as UserService

router = APIRouter(tags=["reports"])


async def require_admin(user: Dict[str, Any] = Depends(get_current_user)):
    if not user.get("is_admin"):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Admin required")
    return user


@router.get("/certificate/{user_id}")
def generate_certificate(user_id: str, user: Dict[str, Any] = Depends(require_admin)):
    target = UserService.find_by_id(user_id)
    if not target:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="User not found")

    doc = Document()
    title = doc.add_heading("Certificate of Achievement", level=0)
    title.alignment = 1

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("This certifies that ").bold = False
    p.add_run(target.username or target.email).bold = True
    p.add_run(
        " has successfully completed cybersecurity awareness training "
        "with Phalanx Cyber Academy."
    )

    doc.add_paragraph()
    doc.add_paragraph(f"Total XP earned: {target.total_xp}")
    doc.add_paragraph(f"Issued on behalf of Phalanx Cyber Academy.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{target.username or 'user'}_certificate.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
