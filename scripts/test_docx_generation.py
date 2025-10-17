#!/usr/bin/env python3
"""
Test script to verify DOCX generation works correctly
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

def test_docx_creation():
    """Test basic DOCX creation to verify python-docx is working"""
    try:
        # Create a simple document
        doc = Document()
        
        # Add a heading
        doc.add_heading('5.1 Test Plan and Results', level=2)
        heading_para = doc.add_paragraph('Test Plan No: STP-001-01')
        heading_para.runs[0].bold = True
        
        # Create a simple table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.5)
        
        # Row 1: Screen Design Ref No
        row_cells = table.rows[0].cells
        row_cells[0].text = 'Screen Design Ref No'
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = 'Figure 4.1: Login Page'
        
        # Row 2: Description / Scenario
        row_cells = table.rows[1].cells
        row_cells[0].text = 'Description / Scenario'
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = 'First Page seen by User by typing in the URL and logging in'
        
        # Row 3: Expected Results
        row_cells = table.rows[2].cells
        row_cells[0].text = 'Expected Results'
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = 'Display the Login page and logging in'
        
        # Row 4: Procedure
        row_cells = table.rows[3].cells
        row_cells[0].text = 'Procedure:'
        row_cells[0].paragraphs[0].runs[0].bold = True
        merged_cell = row_cells[0].merge(row_cells[1])
        
        # Clear and add procedure
        merged_cell.text = ''
        proc_para = merged_cell.paragraphs[0]
        proc_run = proc_para.add_run('Procedure:')
        proc_run.bold = True
        
        # Add procedure steps
        steps = [
            'Open Google Chrome',
            'Type in URL on browser: http://bmis.000.pe/login',
            'Display login page',
            'Enter Username "admin1@test.com" and password "Testing123"'
        ]
        
        for step in steps:
            step_para = merged_cell.add_paragraph(step)
            step_para.style = 'List Number'
        
        # Row 5: Remarks
        row_cells = table.rows[4].cells
        row_cells[0].text = 'Remarks'
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = 'Passed'
        
        # Save to test file
        output_path = project_root / 'test-docx-output.docx'
        doc.save(str(output_path))
        
        print(f"✓ Successfully created test DOCX file: {output_path}")
        print(f"✓ File size: {output_path.stat().st_size} bytes")
        print("✓ DOCX generation is working correctly!")
        
        # Clean up
        output_path.unlink()
        print("✓ Cleanup complete")
        
        return True
        
    except Exception as e:
        print(f"✗ Error testing DOCX creation: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == '__main__':
    success = test_docx_creation()
    sys.exit(0 if success else 1)
